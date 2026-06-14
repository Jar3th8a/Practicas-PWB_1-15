from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "evidencias" / "Reporte_Practica_14_Optimizacion_Rendimiento.docx"
CAPTURES = ROOT / "evidencias" / "capturas"

TITLE = "Práctica 14\nOptimización y Rendimiento"
SUBTITLE = "Caché, lazy loading, N+1 y code splitting"
STUDENT = "Alejandro Avalos Espinosa"
DATE_TXT = "11 de junio de 2026"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 24, 31)
MUTED = RGBColor(92, 101, 112)
LIGHT_FILL = "EEF4FB"
GRID_FILL = "F6F8FB"


def set_run_font(run, *, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_table_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Title", 24, INK),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=12 if level == 1 else 8, after=6)
    r = p.add_run(text)
    set_run_font(r, bold=True)
    return p


def add_body(doc, text, *, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, after=6)
    r = p.add_run(text)
    set_run_font(r, italic=italic, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=3)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=10)
    r = p.add_run(text)
    set_run_font(r, size=10, italic=True, color=MUTED)
    return p


def add_figure(doc, number, title, image_name, description, caption_text):
    add_heading(doc, f"Evidencia {number}. {title}", level=2)
    add_body(doc, description)
    image_path = CAPTURES / image_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=3, after=0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.35))
    add_caption(doc, caption_text)


def add_cover(doc):
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=90, after=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=4)
    r = p.add_run("Reporte de práctica")
    set_run_font(r, size=18, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=2)
    r = p.add_run(TITLE)
    set_run_font(r, size=27, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=18)
    r = p.add_run(SUBTITLE)
    set_run_font(r, size=12, italic=True, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths = [1.55, 4.55]
    metadata = [
        ("Alumno", STUDENT),
        ("Asignatura", "Desarrollo Web Full-Stack"),
        ("Tema", "Optimización del rendimiento en una SPA con Laravel y Vue"),
        ("Fecha", DATE_TXT),
    ]

    for idx, (label, value) in enumerate(metadata):
        left, right = table.rows[idx].cells
        set_table_cell_width(left, widths[0])
        set_table_cell_width(right, widths[1])
        shade_cell(left, LIGHT_FILL)

        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(lp, after=0)
        lr = lp.add_run(label)
        set_run_font(lr, bold=True, color=DARK_BLUE)

        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(rp, after=0)
        rr = rp.add_run(value)
        set_run_font(rr, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=16, after=0)
    r = p.add_run("Documento de entrega académica")
    set_run_font(r, size=11, italic=True, color=MUTED)


def add_summary_table(doc):
    add_heading(doc, "2. Resumen de optimizaciones", level=1)
    add_body(
        doc,
        "La práctica se trabajó sobre la base del proyecto integrador para concentrar los cambios de rendimiento en una sola experiencia de navegación, con mejoras en servidor, caché y carga diferida del frontend.",
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    headers = ["Área", "Ajuste realizado"]
    row = table.rows[0].cells
    widths = [1.85, 4.65]
    for i, text in enumerate(headers):
        set_table_cell_width(row[i], widths[i])
        shade_cell(row[i], GRID_FILL)
        p = row[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, after=0)
        r = p.add_run(text)
        set_run_font(r, bold=True, color=DARK_BLUE)

    rows = [
        ("N+1", "Carga anticipada de relaciones con eager loading y relación global en Producto."),
        ("Caché", "Versionado de claves para productos y categorías, con invalidación al crear, editar o eliminar."),
        ("Índices", "Índices en categoria_id, precio y un índice compuesto sobre nombre + precio."),
        ("Frontend", "Rutas y componentes cargados de forma asíncrona para reducir el bundle inicial."),
        ("Medición", "Bundle visualizado con rollup-plugin-visualizer y verificación de endpoints JSON."),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        set_table_cell_width(cells[0], widths[0])
        set_table_cell_width(cells[1], widths[1])

        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p0, after=0)
        r0 = p0.add_run(label)
        set_run_font(r0, bold=True, color=INK)

        p1 = cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p1, after=0)
        r1 = p1.add_run(value)
        set_run_font(r1, color=INK)


def build_report():
    doc = Document()
    style_document(doc)

    add_cover(doc)
    doc.add_page_break()

    add_heading(doc, "1. Objetivo", level=1)
    add_body(
        doc,
        "Identificar y corregir cuellos de botella comunes en una aplicación con Laravel y Vue, enfocándose en consultas N+1, cacheo de datos de consulta frecuente, reducción del tamaño inicial del frontend mediante lazy loading y medición de resultados con herramientas de análisis.",
    )

    add_summary_table(doc)

    add_heading(doc, "3. Desarrollo realizado", level=1)
    add_bullet(doc, "Se instaló y configuró Laravel Telescope para revisar solicitudes, consultas y eventos de la aplicación.")
    add_bullet(doc, "Se aplicó eager loading en el modelo y en los controladores para evitar cargas repetidas de relaciones.")
    add_bullet(doc, "Se implementó una capa de caché con claves versionadas y limpieza automática cuando cambia el catálogo.")
    add_bullet(doc, "Se añadieron índices de base de datos para mejorar filtros y búsquedas frecuentes sobre productos.")
    add_bullet(doc, "Se migraron rutas y componentes de Vue a carga asíncrona para reducir el bundle inicial.")
    add_bullet(doc, "Se generó un visualizador de bundle con Vite para revisar el peso relativo de cada dependencia.")

    add_heading(doc, "4. Evidencias", level=1)

    evidence = [
        (
            1,
            "Catálogo cargado",
            "practica14_01_catalogo.png",
            "La vista principal carga productos desde la API a través del proxy de desarrollo de Vite, mostrando tarjetas, precios y acciones sin errores de CORS.",
            "Figura 1. Catálogo inicial con productos recuperados desde la API.",
        ),
        (
            2,
            "Respuesta JSON de la API",
            "practica14_02_api_productos_json.png",
            "El endpoint devuelve la colección paginada con estructura JSON correcta, lo que confirma la comunicación entre el frontend y Laravel.",
            "Figura 2. Respuesta JSON del endpoint /api/productos.",
        ),
        (
            3,
            "Visualizador del bundle",
            "practica14_03_bundle_visualizer.png",
            "El gráfico de treemap permite identificar con claridad qué dependencias aportan más peso al bundle inicial y qué módulos se cargan de forma diferida.",
            "Figura 3. Análisis visual del bundle obtenido con Vite.",
        ),
        (
            4,
            "Rutas de Telescope",
            "practica14_04_telescope_routes.png",
            "La salida de Artisan confirma que Telescope quedó instalado con sus rutas internas para revisar requests, queries, modelos y otros eventos del sistema.",
            "Figura 4. Rutas registradas por Laravel Telescope.",
        ),
    ]

    for idx, title, image_name, description, caption in evidence:
        add_figure(doc, idx, title, image_name, description, caption)
        if idx != len(evidence):
            doc.add_page_break()

    doc.add_page_break()
    add_heading(doc, "5. Resultado final", level=1)
    add_body(
        doc,
        "La aplicación quedó lista para trabajar con un flujo más eficiente en servidor y cliente: las consultas al catálogo se responden con mejor reutilización de datos, el frontend se carga de forma más ligera y el análisis del bundle facilita detectar dependencias pesadas.",
    )
    add_body(
        doc,
        "Como cierre, la práctica deja una base funcional para seguir midiendo mejoras de rendimiento sin perder claridad en el código ni en la experiencia de uso.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
