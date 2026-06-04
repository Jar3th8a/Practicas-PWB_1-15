from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\alesh\practica1-autenticacion\practica7")
OUT = ROOT / "evidencias" / "Reporte_Practica_07_Paginacion_Filtros.docx"
CAPTIONS = ROOT / "evidencias" / "capturas"

TITLE = "Practica 07\nPaginacion y Filtros"
SUBTITLE = "Laravel + Vue.js con busqueda, categoria, rango de precio y paginacion server-side"
STUDENT = "Alejandro Avalos Espinosa"
DATE_TXT = "2 de junio de 2026"


def set_run_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_image_name(run, name):
    drawing = run._r.drawing_lst[0]
    ns = {
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    }
    doc_pr = drawing.find(".//wp:docPr", namespaces=ns)
    if doc_pr is not None:
        doc_pr.set("name", name)
    c_nv_pr = drawing.find(".//pic:cNvPr", namespaces=ns)
    if c_nv_pr is not None:
        c_nv_pr.set("name", name)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=3)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=3, after=9)
    r = p.add_run(text)
    set_run_font(r, size=10, italic=True, color=RGBColor(80, 80, 80))
    return p


def add_evidence(doc, number, title, image_name, description, caption_text):
    add_heading(doc, f"Evidencia {number}. {title}", level=2)
    add_body(doc, description)
    img = CAPTIONS / image_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=3, after=0)
    r = p.add_run()
    r.add_picture(str(img), width=Inches(6.1))
    set_image_name(r, caption_text)
    add_caption(doc, caption_text)


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Title", 24, RGBColor(31, 78, 121)),
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=120, after=6)
    r = p.add_run("Reporte de practica\n")
    set_run_font(r, size=18, bold=False, color=RGBColor(31, 78, 121))
    r2 = p.add_run(TITLE)
    set_run_font(r2, size=28, bold=True, color=RGBColor(31, 78, 121))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=16)
    r = p.add_run(SUBTITLE)
    set_run_font(r, size=12, italic=True, color=RGBColor(80, 80, 80))

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths = [1.55, 4.55]
    rows = [
        ("Alumno", STUDENT),
        ("Asignatura", "Desarrollo Web Full-Stack"),
        ("Practica", "07 - Paginacion y filtros avanzados"),
        ("Tecnologias", "Laravel API, Vue 3, Vue Router, Axios"),
        ("Fecha", DATE_TXT),
    ]
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        set_cell_width(c0, widths[0])
        set_cell_width(c1, widths[1])
        shade_cell(c0, "EAF1FB")
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p0, after=0)
        r0 = p0.add_run(label)
        set_run_font(r0, bold=True)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p1, after=0)
        r1 = p1.add_run(value)
        set_run_font(r1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=18, after=0)
    r = p.add_run("Documento preparado para entrega academica.")
    set_run_font(r, size=11, italic=True, color=RGBColor(80, 80, 80))


def main():
    doc = Document()
    configure_document(doc)

    add_title_page(doc)
    doc.add_page_break()

    add_heading(doc, "1. Objetivo", level=1)
    add_body(
        doc,
        "Desarrollar un catalogo de productos con paginacion del lado del servidor, busqueda por texto, filtrado por categoria y rango de precio, sincronizando el estado de los filtros con la URL para que la experiencia sea compartible y consistente.",
    )

    add_heading(doc, "2. Desarrollo realizado", level=1)
    add_bullet(doc, "Se creo un backend Laravel con endpoints REST para productos y categorias.")
    add_bullet(doc, "Se implementaron scopes Eloquent para busqueda, categoria y rango de precio.")
    add_bullet(doc, "Se desarrollo un frontend Vue 3 con composable de filtros, panel lateral y navegacion paginada.")
    add_bullet(doc, "Se verifico la integracion Axios + Laravel mediante pruebas de navegacion y JSON en API.")

    add_heading(doc, "3. Evidencias", level=1)
    evidence = [
        (1, "Catalogo inicial", "practica7_01_catalogo_inicial.png", "Vista principal del catalogo con filtros laterales, tarjetas de productos y controles de paginacion visibles al abrir la aplicacion.", "Figura 1. Catalogo inicial de productos."),
        (2, "Busqueda por nombre", "practica7_02_busqueda_laptop.png", "El campo de busqueda contiene el termino 'Laptop' y el listado se reduce a los resultados relacionados.", "Figura 2. Resultado de busqueda por texto."),
        (3, "Filtro por categoria", "practica7_03_categoria_hogar.png", "Se selecciono la categoria 'Hogar' desde el panel lateral para validar el filtrado por categoria.", "Figura 3. Aplicacion de filtro por categoria."),
        (4, "Rango de precios", "practica7_04_rango_precios.png", "Se aplico un rango de precios entre 1000 y 1200 para comprobar que el catalogo filtra por limites minimo y maximo.", "Figura 4. Aplicacion de filtro por rango de precio."),
        (5, "Paginacion", "practica7_05_paginacion_pagina2.png", "Se navego a la pagina 2 del catalogo para mostrar la paginacion server-side y el conteo total de paginas.", "Figura 5. Navegacion a la pagina 2."),
        (6, "Detalle de producto", "practica7_06_detalle_producto.png", "Se abrio el detalle del producto seleccionado para verificar la ruta dinamica y la categoria asociada.", "Figura 6. Vista detallada del producto."),
        (7, "Respuesta JSON de API", "practica7_07_api_json.png", "Se consulto el endpoint REST con filtros en la URL y se documento la respuesta JSON del backend.", "Figura 7. Respuesta JSON del endpoint de productos."),
    ]
    for idx, title, img, desc, caption in evidence:
        add_evidence(doc, idx, title, img, desc, caption)
        if idx != len(evidence):
            doc.add_page_break()

    doc.add_page_break()
    add_heading(doc, "4. Resultado final", level=1)
    add_bullet(doc, "El catalogo responde con filtros reales y paginacion desde Laravel.")
    add_bullet(doc, "La URL refleja el estado de busqueda, categoria y paginacion.")
    add_bullet(doc, "Se incluyeron evidencias visuales y un reporte Word para entrega academica.")
    add_bullet(doc, "Validaciones ejecutadas: php artisan test y npm run build.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
