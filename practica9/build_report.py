from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\alesh\practica1-autenticacion\practica9")
OUT = ROOT / "evidencias" / "Reporte_Practica_09_Validaciones_Avanzadas.docx"
CAPTIONS = ROOT / "evidencias" / "capturas"

TITLE = "Practica 09\nValidaciones Avanzadas"
SUBTITLE = "Laravel Form Requests en backend y VeeValidate + Yup en frontend"
STUDENT = "Alejandro Avalos Espinosa"
DATE_TXT = "2 de junio de 2026"


def set_run_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_image_name(run, name):
    drawing = run._r.drawing_lst[0]
    ns = {
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    }
    doc_pr = drawing.find(".//wp:docPr", namespaces=ns)
    if doc_pr is not None:
        doc_pr.set("name", name)
    c_nv_pr = drawing.find(".//pic:cNvPr", namespaces=ns)
    if c_nv_pr is not None:
        c_nv_pr.set("name", name)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=3)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=3, after=9)
    r = p.add_run(text)
    set_run_font(r, size=10, italic=True, color=RGBColor(80, 80, 80))
    return p


def add_evidence(doc, number, title, image_name, description, caption_text):
    add_heading(doc, f"Evidencia {number}. {title}", level=2)
    add_body(doc, description)
    img = CAPTIONS / image_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=3, after=0)
    r = p.add_run()
    r.add_picture(str(img), width=Inches(6.1))
    set_image_name(r, caption_text)
    add_caption(doc, caption_text)


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Title", 24, RGBColor(31, 78, 121)),
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=120, after=6)
    r = p.add_run("Reporte de practica\n")
    set_run_font(r, size=18, bold=False, color=RGBColor(31, 78, 121))
    r2 = p.add_run(TITLE)
    set_run_font(r2, size=28, bold=True, color=RGBColor(31, 78, 121))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=16)
    r = p.add_run(SUBTITLE)
    set_run_font(r, size=12, italic=True, color=RGBColor(80, 80, 80))

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths = [1.55, 4.55]
    rows = [
        ("Alumno", STUDENT),
        ("Asignatura", "Desarrollo Web Full-Stack"),
        ("Practica", "09 - Validaciones avanzadas"),
        ("Tecnologias", "Laravel Form Requests, VeeValidate, Yup, Vue 3, Axios"),
        ("Fecha", DATE_TXT),
    ]
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        set_cell_width(c0, widths[0])
        set_cell_width(c1, widths[1])
        shade_cell(c0, "EAF1FB")
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p0, after=0)
        r0 = p0.add_run(label)
        set_run_font(r0, bold=True)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p1, after=0)
        r1 = p1.add_run(value)
        set_run_font(r1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=18, after=0)
    r = p.add_run("Documento preparado para entrega academica.")
    set_run_font(r, size=11, italic=True, color=RGBColor(80, 80, 80))


def main():
    doc = Document()
    configure_document(doc)

    add_title_page(doc)
    doc.add_page_break()

    add_heading(doc, "1. Objetivo", level=1)
    add_body(
        doc,
        "Implementar validaciones robustas en dos capas: Laravel Form Requests en el backend, con mensajes en español para cada campo, y VeeValidate con Yup en el frontend para mostrar errores en tiempo real y reforzar la calidad de los datos antes de enviarlos al servidor.",
    )

    add_heading(doc, "2. Desarrollo realizado", level=1)
    add_bullet(doc, "Se crearon StoreProductoRequest y UpdateProductoRequest con reglas de validacion y mensajes en español.")
    add_bullet(doc, "Se mantuvo la validacion de backend para nombre, descripcion, precio, stock, categoria e imagen.")
    add_bullet(doc, "Se implemento VeeValidate con Yup para validar el formulario de productos antes del envio.")
    add_bullet(doc, "Se reutilizo un componente InputField.vue para unificar label, campo y mensajes de error.")

    add_heading(doc, "3. Evidencias", level=1)
    evidence = [
        (1, "Formulario con errores", "practica9_01_frontend_errores.png", "Se dejo el formulario en blanco para comprobar que VeeValidate muestra los mensajes de error en tiempo real en nombre, precio y stock.", "Figura 1. Errores visibles en el formulario del frontend."),
        (2, "Creacion exitosa", "practica9_02_creacion_exitosa.png", "Se completo el formulario con datos validos y se confirmo que el producto se crea correctamente sin recargar la pagina.", "Figura 2. Registro exitoso de un producto."),
        (3, "Edicion exitosa", "practica9_03_edicion_exitosa.png", "Se abrio un producto existente, se modificaron sus datos y se guardo con exito mediante el mismo formulario reutilizable.", "Figura 3. Actualizacion correcta de un producto."),
        (4, "Vista previa de imagen", "practica9_04_preview_imagen.png", "Al seleccionar una imagen, el formulario muestra una vista previa local antes de enviar la informacion al servidor.", "Figura 4. Vista previa de imagen seleccionada."),
        (5, "Error 422 del backend", "practica9_05_error_backend_422.png", "Se envio informacion incompleta para provocar la respuesta de validacion de Laravel con mensajes en español.", "Figura 5. Respuesta de validacion del servidor."),
        (6, "JSON 422 de API", "practica9_06_api_422_json.png", "Se consulto el endpoint de productos con datos invalidos para documentar el JSON de error que retorna el backend.", "Figura 6. Respuesta JSON de validacion."),
    ]
    for idx, title, img, desc, caption in evidence:
        add_evidence(doc, idx, title, img, desc, caption)
        if idx != len(evidence):
            doc.add_page_break()

    doc.add_page_break()
    add_heading(doc, "4. Resultado final", level=1)
    add_bullet(doc, "El backend valida todos los campos relevantes mediante Form Requests y mensajes en español.")
    add_bullet(doc, "La interfaz Vue muestra errores en tiempo real con VeeValidate + Yup y conserva la validacion del servidor.")
    add_bullet(doc, "El formulario reutilizable InputField.vue centraliza la presentacion de los campos y los mensajes de error.")
    add_bullet(doc, "Validaciones ejecutadas: php artisan migrate:fresh --seed, php artisan test y npm run build.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

