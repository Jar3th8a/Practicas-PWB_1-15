from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "evidencias" / "Reporte_Practica_15_Despliegue_Produccion.docx"
CAPTURES = ROOT / "evidencias" / "capturas"

TITLE = "Práctica 15\nDespliegue en Producción"
SUBTITLE = "Llevar la aplicación Full-Stack a un servidor real"
STUDENT = "Alejandro Avalos Espinosa"
DATE_TXT = "11 de junio de 2026"

INK = RGBColor(21, 28, 38)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 101, 112)
SOFT_FILL = "EEF4FB"


def set_run_font(run, *, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Title", 24, INK),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=12 if level == 1 else 8, after=6)
    r = p.add_run(text)
    set_run_font(r, bold=True)
    return p


def add_body(doc, text, *, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, after=6)
    r = p.add_run(text)
    set_run_font(r, italic=italic, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=3)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=10)
    r = p.add_run(text)
    set_run_font(r, size=10, italic=True, color=MUTED)
    return p


def add_figure(doc, number, title, image_name, description, caption_text):
    add_heading(doc, f"Evidencia {number}. {title}", level=2)
    add_body(doc, description)
    img = CAPTURES / image_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=3, after=0)
    r = p.add_run()
    r.add_picture(str(img), width=Inches(6.35))
    add_caption(doc, caption_text)


def add_cover(doc):
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=92, after=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=4)
    r = p.add_run("Reporte de práctica")
    set_run_font(r, size=18, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=2)
    r = p.add_run(TITLE)
    set_run_font(r, size=27, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=18)
    r = p.add_run(SUBTITLE)
    set_run_font(r, size=12, italic=True, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths = [1.55, 4.55]
    meta = [
        ("Alumno", STUDENT),
        ("Asignatura", "Desarrollo Web Full-Stack"),
        ("Tema", "Despliegue y operación en producción"),
        ("Fecha", DATE_TXT),
    ]

    for idx, (label, value) in enumerate(meta):
        c0, c1 = table.rows[idx].cells
        set_cell_width(c0, widths[0])
        set_cell_width(c1, widths[1])
        shade_cell(c0, SOFT_FILL)

        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p0, after=0)
        r0 = p0.add_run(label)
        set_run_font(r0, bold=True, color=DARK_BLUE)

        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p1, after=0)
        r1 = p1.add_run(value)
        set_run_font(r1, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=16, after=0)
    r = p.add_run("Documento de entrega académica")
    set_run_font(r, size=11, italic=True, color=MUTED)


def add_summary(doc):
    add_heading(doc, "2. Objetivo", level=1)
    add_body(
        doc,
        "Desplegar la aplicación completa Vue + Laravel en un entorno de producción, documentando la preparación del servidor, el uso de variables de entorno seguras, la publicación de la SPA con Nginx, la automatización del despliegue con GitHub Actions y el mantenimiento del worker de colas con Supervisor.",
    )

    add_heading(doc, "3. Resumen técnico", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = ["Elemento", "Propósito"]
    widths = [2.0, 4.4]
    head_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_width(head_cells[i], widths[i])
        shade_cell(head_cells[i], SOFT_FILL)
        p = head_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, after=0)
        r = p.add_run(text)
        set_run_font(r, bold=True, color=DARK_BLUE)

    rows = [
        ("Nginx", "Servir la SPA y enrutar /api hacia Laravel en producción."),
        ("Laravel", "Ejecutar migraciones, cachear configuración y atender la API."),
        ("Vue", "Compilar la interfaz y publicar los archivos estáticos dentro de /dist."),
        ("Supervisor", "Mantener el worker de colas y el proceso de Reverb activos."),
        ("GitHub Actions", "Automatizar el despliegue en cada push a main."),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], widths[0])
        set_cell_width(cells[1], widths[1])
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p0, after=0)
        r0 = p0.add_run(label)
        set_run_font(r0, bold=True, color=INK)
        p1 = cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p1, after=0)
        r1 = p1.add_run(value)
        set_run_font(r1, color=INK)


def build_report():
    doc = Document()
    style_document(doc)

    add_cover(doc)
    doc.add_page_break()
    add_summary(doc)

    add_heading(doc, "4. Desarrollo realizado", level=1)
    add_bullet(doc, "Se prepararon archivos de entorno para backend y frontend con valores pensados para producción.")
    add_bullet(doc, "Se definió una configuración de Nginx para servir la SPA y el backend detrás de PHP-FPM.")
    add_bullet(doc, "Se documentó un servicio de Supervisor para el worker de colas y el servidor de Reverb.")
    add_bullet(doc, "Se creó un flujo de GitHub Actions para desplegar backend y frontend por SSH.")
    add_bullet(doc, "Se ejecutó el build de Vue y se aplicó la optimización de Laravel para verificar el cierre operativo.")

    add_heading(doc, "5. Evidencias", level=1)

    evidences = [
        (
            1,
            "Variables de entorno del backend",
            "practica15_01_env_backend.png",
            "El archivo de producción del backend concentra los parámetros del entorno real: base de datos, Redis, Reverb y modo production.",
            "Figura 1. Variables de entorno para Laravel en producción.",
        ),
        (
            2,
            "Variables de entorno del frontend",
            "practica15_02_env_frontend.png",
            "El frontend toma la URL pública de la API y conserva la versión y los datos de Reverb necesarios para la comunicación en vivo.",
            "Figura 2. Variables de entorno para Vue en producción.",
        ),
        (
            3,
            "Configuración de Nginx",
            "practica15_03_nginx.png",
            "La configuración define la raíz de la SPA, la ruta de la API y el acceso a los archivos de storage mediante alias.",
            "Figura 3. Virtual host de Nginx para producción.",
        ),
        (
            4,
            "Configuración de Supervisor",
            "practica15_04_supervisor.png",
            "El archivo mantiene el worker de colas y el servidor de Reverb activos, con reinicio automático en caso de fallas.",
            "Figura 4. Supervisor para procesos de cola y Reverb.",
        ),
        (
            5,
            "Workflow de despliegue",
            "practica15_05_workflow.png",
            "El pipeline de GitHub Actions ejecuta pull, instala dependencias, aplica migraciones, limpia cachés y construye el frontend.",
            "Figura 5. Pipeline de despliegue automático con GitHub Actions.",
        ),
        (
            6,
            "Build del frontend",
            "practica15_06_build_output.png",
            "La salida de Vite confirma que la compilación de producción concluye con éxito y muestra el tamaño de los bundles resultantes.",
            "Figura 6. Resultado del build de producción de Vue.",
        ),
        (
            7,
            "Optimización de Laravel",
            "practica15_07_optimize_output.png",
            "El comando optimize cachea configuración, rutas y vistas para dejar Laravel listo para servir con menor sobrecarga.",
            "Figura 7. Optimización final de Laravel para producción.",
        ),
    ]

    for idx, title, image_name, description, caption in evidences:
        add_figure(doc, idx, title, image_name, description, caption)
        if idx != len(evidences):
            doc.add_page_break()

    doc.add_page_break()
    add_heading(doc, "6. Resultado final", level=1)
    add_body(
        doc,
        "La práctica deja documentada una ruta de despliegue completa para la aplicación Full-Stack: entorno de producción preparado, servidor web configurado, colas supervisadas, publicación automática y verificación de build y cachés.",
    )
    add_body(
        doc,
        "Con esta base, el proyecto puede moverse a un VPS real siguiendo el mismo esquema, reduciendo tiempos manuales y manteniendo consistencia entre cada entrega.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
